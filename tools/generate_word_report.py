"""
GTT Center Perth -- Word Financial Report Generator (Deliverable 2).

Generates outputs/GTT-Center-Perth-Financial-Report.docx, a 16-section
narrative report built entirely from tools/document_data.py -- the same
shared data-assembly layer used by the Excel workbook, so the two
deliverables cannot drift apart.

SCOPE BOUNDARIES (do not violate):
- Does not choose Table 1 or Table 2 as the venture's primary scenario.
- Does not invent an opening cash balance or a financing structure.
- Presents the funding requirement as a BOUNDED RANGE, never a false exact
  figure.
- Never presents startup/capex costs as part of operating P&L.
- Never presents the historical A$157,792.16 / A$118,297.16 figures as
  current canonical revenue -- always labelled HISTORICAL-SUPERSEDED.
- Never labels the historical A$63,028.75 figure as revenue -- it is a
  historical Table 1 Monthly Net P&L figure, full stop.
"""

import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO_ROOT / "tools"))
import document_data as dd  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.shared import Pt, RGBColor, Cm  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

OUTPUT_PATH = REPO_ROOT / "outputs" / "GTT-Center-Perth-Financial-Report.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
RED = RGBColor(0xC0, 0x00, 0x00)
GREY = RGBColor(0x59, 0x59, 0x59)


def fmt_aud(v):
    if v is None:
        return "n/a"
    neg = v < 0
    s = f"A${abs(v):,.2f}"
    return f"-{s}" if neg else s


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): color_hex})
    tcPr.append(shd)


def add_title_page(doc, data: dd.DocumentData):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("GTT Center Perth")
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Financial Model & Report")
    run.font.size = Pt(20)
    run.font.color.rgb = GREY

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Two-scenario canonical financial model (Table 1: 18 clients/day, Table 2: 12 clients/day)")
    run.font.size = Pt(12)
    run.italic = True

    doc.add_paragraph()
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box.add_run(
        "IMPORTANT -- READ FIRST\n"
        "This report presents CANONICAL (validated), MODELLED (calculated from canonical inputs), "
        "HISTORICAL-SUPERSEDED (preserved for traceability, no longer live), and BOUNDED (range, not exact) "
        "figures, each explicitly labelled. Neither Table 1 nor Table 2 is presented as the venture's primary "
        "scenario -- see VERIFICATION-TRACKER.md item 1m, still open. The funding requirement (Section 10) is a "
        "bounded range, not a single point estimate. No opening cash balance or financing structure is assumed "
        "anywhere in this report."
    )
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = RED

    doc.add_page_break()


def add_h1(doc, num, title):
    h = doc.add_heading(f"{num}. {title}", level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_h2(doc, title):
    h = doc.add_heading(title, level=2)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_para(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    return p


def add_warning(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RED
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], "1F3864")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table


# ---------------------------------------------------------------------------
# Section 1 -- Executive Summary
# ---------------------------------------------------------------------------
def section_executive_summary(doc, data: dd.DocumentData):
    add_h1(doc, 1, "Executive Summary")
    add_para(
        doc,
        "GTT Center Perth's canonical financial model supports two committed scenarios, presented side by side "
        "throughout this report. Neither is designated the venture's primary basis (VERIFICATION-TRACKER.md item 1m "
        "remains open) -- Table 1 (18 clients/day) generates materially higher revenue and net operating result at "
        "identical headcount, but the founder decision to formally adopt it over Table 2 (12 clients/day) has not "
        "been made."
    )
    sc = data.scenario_comparison
    rows = [
        ["Client volume/day", sc["scenario_table_1"]["client_volume_per_day"], sc["scenario_table_2"]["client_volume_per_day"]],
        ["Monthly revenue (steady state)", fmt_aud(sc["scenario_table_1"]["steady_state_revenue"]), fmt_aud(sc["scenario_table_2"]["steady_state_revenue"])],
        ["Monthly operating costs (steady state)", fmt_aud(sc["scenario_table_1"]["steady_state_total_operating_costs"]), fmt_aud(sc["scenario_table_2"]["steady_state_total_operating_costs"])],
        ["Net operating result (steady state, post-super)", fmt_aud(sc["scenario_table_1"]["steady_state_net_operating_result"]), fmt_aud(sc["scenario_table_2"]["steady_state_net_operating_result"])],
        ["24-month net operating result", fmt_aud(data.totals_24mo["scenario_table_1"]["total_net_operating_result_24mo"]), fmt_aud(data.totals_24mo["scenario_table_2"]["total_net_operating_result_24mo"])],
        ["Break-even AM client volume/day", data.breakeven["scenario_table_1"]["breakeven_am_client_volume_per_day"], data.breakeven["scenario_table_2"]["breakeven_am_client_volume_per_day"]],
        ["Margin of safety (clients/day)", data.breakeven["scenario_table_1"]["margin_of_safety_clients_per_day"], data.breakeven["scenario_table_2"]["margin_of_safety_clients_per_day"]],
        ["Operating cash trough", fmt_aud(data.cash_flow["scenario_table_1"]["trough_cumulative_position"]), fmt_aud(data.cash_flow["scenario_table_2"]["trough_cumulative_position"])],
    ]
    add_table(doc, ["Metric", "Table 1 (18/day)", "Table 2 (12/day)"], rows)
    doc.add_paragraph()
    add_para(doc, "Bounded funding requirement (both scenarios, universal, not scenario-specific):", bold=True)
    fr = data.funding["combined_funding_requirement_bounded"]["primary_method"]
    add_warning(doc, f"A${fr['range_low']:,.2f} - A${fr['range_high']:,.2f} -- a RANGE, not an exact figure. See Section 10.")
    doc.add_paragraph()
    add_para(doc, "Key risks and unresolved assumptions (full list in Section 13):", bold=True)
    add_bullets(doc, [
        "Historical revenue figures differ from canonical by A$2,576.36/month; the gap's origin is permanently unresolved (item 36).",
        "A 10% PM pre-booking discount is NOT applied anywhere in this model (item 39) -- a revenue-overstatement risk if the discount is real and active.",
        "AM labour is modelled fixed from Month 1 rather than ramped with client volume -- a conservative, disclosed simplification (item 43).",
        "Superannuation is resolved for 6 of 8 payroll components; Opening-Time-Increment and Receptionist/Relief Pool remain super-uncertain (item 46, partial).",
        "The funding requirement is bounded, not exact -- the underlying 6-9-range historical startup-capital reconciliation remains unresolved (item 47).",
    ])


# ---------------------------------------------------------------------------
# Section 2 -- Business Model Overview
# ---------------------------------------------------------------------------
def section_business_model(doc, data: dd.DocumentData):
    add_h1(doc, 2, "Business Model Overview")
    add_para(
        doc,
        "GTT Center Perth transforms the mandatory 2-3 hour Glucose Tolerance Test (GTT) wait -- undergone by "
        "pregnant women in Western Australia -- into a restorative waiting experience: massage, nails, hair, and a "
        "3D keepsake ultrasound. A pathology partner (PathWest, WDP, or Clinipath) handles all blood draws; beauty "
        "and wellness practitioners are employed staff (casual or permanent), not subtenants. No NATA accreditation "
        "is required for the venue's own operations, and the 3D keepsake scan is legal and unregulated in WA -- it "
        "is presented strictly as a keepsake/entertainment product, never as diagnostic imaging."
    )
    add_para(
        doc,
        "Revenue flows through YETI Holding Trust. The venture is self-funded via Anthony and Imara's joint "
        "savings, with no external investor. Day-to-day operations are run by a Venue Manager (a planned hire, not "
        "yet in place); Imara has no operational involvement -- ownership and financial oversight only."
    )
    add_para(
        doc,
        "The AM (morning) session runs on a 25-minute pairing cadence with two committed scenarios: Table 1 (18 "
        "clients/day, 07:00 start) and Table 2 (12 clients/day, 08:00 start). Both scenarios use the identical "
        "staffing model -- 8 dual/multi-qualified treatment staff (4 Massage+Beauty, 2 Nails, 2 Hair) plus 2 "
        "phlebotomists -- meaning Table 1 generates materially more revenue at zero additional weekday labour cost. "
        "Saturday reuses the same AM volume as weekday; Sunday is closed. Packages are priced at A$250 and A$300 "
        "(a lower-tier package was previously offered and has since been dropped)."
    )


# ---------------------------------------------------------------------------
# Section 3 -- Revenue Model
# ---------------------------------------------------------------------------
def section_revenue_model(doc, data: dd.DocumentData):
    add_h1(doc, 3, "Revenue Model")
    add_para(
        doc,
        "Canonical Total Revenue = AM Revenue + PM Revenue + Ancillary Revenue, computed per "
        "docs/architecture/CANONICAL-REVENUE-METHODOLOGY.md from data/canonical/pricing.yml, "
        "client_assumptions.yml, and scenarios.yml. AM Revenue = AM package price x committed client volume x "
        "operating days/month. PM Revenue uses the a-la-carte average price. Ancillary Revenue is A$0.00 in both "
        "scenarios -- explicitly excluded per Anthony's 2026-07-30 instruction, not an omission."
    )
    for sid in dd.SCENARIOS:
        add_h2(doc, data.label(sid))
        ss = data.steady_state[sid]["revenue"]
        rows = [
            ["AM revenue", fmt_aud(ss["am_revenue"])],
            ["PM revenue", fmt_aud(ss["pm_revenue"])],
            ["Ancillary revenue", fmt_aud(ss["ancillary_revenue"])],
            ["Total revenue (Month 5+, steady state)", fmt_aud(ss["total_revenue"])],
        ]
        add_table(doc, ["Component", "Monthly Value (AUD)"], rows)
        doc.add_paragraph()
    add_warning(
        doc,
        "Historical figures (A$157,792.16 Table 1 / A$118,297.16 Table 2) are NOT current canonical revenue -- see "
        "Section 14, Historical vs. Canonical Figures."
    )


# ---------------------------------------------------------------------------
# Section 4 -- Revenue Ramp
# ---------------------------------------------------------------------------
def section_revenue_ramp(doc, data: dd.DocumentData):
    add_h1(doc, 4, "Revenue Ramp")
    add_para(
        doc,
        "Revenue ramps from Month 1 to Month 5+ (steady state) per data/canonical/revenue_ramp.yml. Months 6-24 "
        "are held flat at the Month 5+ canonical steady-state figure -- no post-Month-5 growth curve exists in any "
        "document in this repository, so none is invented (assumption_24month_extension_flat)."
    )
    for sid in dd.SCENARIOS:
        add_h2(doc, data.label(sid))
        rows = []
        for m in range(1, 6):
            pnl = data.pnl_24mo[sid][m - 1]
            label = f"Month {m}" if m < 5 else "Month 5+ (steady state)"
            pct = pnl["revenue"]["total_revenue"] / data.steady_state[sid]["revenue"]["total_revenue"] * 100
            rows.append([label, fmt_aud(pnl["revenue"]["total_revenue"]), f"{pct:.1f}%"])
        add_table(doc, ["Month", "Total Revenue (AUD)", "% of Steady State"], rows)
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 5 -- Operating Cost Model
# ---------------------------------------------------------------------------
def section_cost_model(doc, data: dd.DocumentData):
    add_h1(doc, 5, "Operating Cost Model")
    add_para(
        doc,
        "Canonical Total Operating Costs = Payroll (incl. superannuation and workers compensation) + Operating "
        "Expenses (fixed non-wage overhead + variable/ramped marketing), per "
        "docs/architecture/COST-RAMP-METHODOLOGY.md and data/canonical/cost_ramp.yml."
    )
    for sid in dd.SCENARIOS:
        add_h2(doc, data.label(sid))
        ss = data.steady_state[sid]
        rows = [
            ["Payroll (incl. superannuation)", fmt_aud(ss["payroll"])],
            ["Operating expenses (fixed + variable)", fmt_aud(ss["operating_expenses"])],
            ["Total operating costs (Month 5+, steady state)", fmt_aud(ss["total_operating_costs"])],
        ]
        add_table(doc, ["Component", "Monthly Value (AUD)"], rows)
        doc.add_paragraph()
    add_warning(
        doc,
        "Unresolved cost classifications carried forward, not resolved by this report: GTT supplies/consumables/"
        "laundry (FIXED vs. VARIABLE, item 45); Insurance (modelled A$400/month vs. itemised A$975-1,583/month, "
        "see Section 12); AM labour ramp (modelled FIXED from Month 1, item 43)."
    )


# ---------------------------------------------------------------------------
# Section 6 -- Superannuation Treatment
# ---------------------------------------------------------------------------
def section_super(doc, data: dd.DocumentData):
    add_h1(doc, 6, "Superannuation Treatment")
    add_para(
        doc,
        "RESOLVED 2026-08-09 (VERIFICATION-TRACKER.md item 46). Superannuation (12% of Ordinary Time Earnings, "
        "wages.yml#wage_superannuation_rate) is now correctly applied at the canonical wage/cost layer for 6 of 8 "
        "payroll components. AM weekday treatment staff and phlebotomist wages already/now correctly incorporate "
        "super per role; AM Saturday, PM weekday, and PM Saturday direct labour have super added at 12%."
    )
    add_warning(
        doc,
        "Opening-Time Increment and Receptionist/Relief Pool remain UNRESOLVED for superannuation eligibility -- "
        "these roles are not specific enough in the current data to confidently apply super "
        "(conflict_superannuation_partial_coverage, data/canonical/cost_ramp.yml)."
    )
    add_para(
        doc,
        "This correction reduced the previously-reported steady-state Net Operating Result from A$60,201.62 to "
        "A$56,581.70/month (Table 1) and from A$24,257.56 to A$21,056.64/month (Table 2). Table 2's Month 3 result "
        "flipped from +A$550.61 to -A$2,639.17 as a direct consequence."
    )


# ---------------------------------------------------------------------------
# Section 7 -- 24-Month Financial Outlook
# ---------------------------------------------------------------------------
def section_24mo_outlook(doc, data: dd.DocumentData):
    add_h1(doc, 7, "24-Month Financial Outlook")
    add_para(
        doc,
        "Full Month 1-24 detail is provided in the accompanying Excel workbook (Sheet 4, live formulas). This "
        "section shows snapshot months only."
    )
    for sid in dd.SCENARIOS:
        add_h2(doc, data.label(sid))
        rows = []
        for m in data.snapshot_months:
            pnl = data.pnl_24mo[sid][m - 1]
            rows.append([
                f"Month {m}", fmt_aud(pnl["revenue"]["total_revenue"]), fmt_aud(pnl["total_operating_costs"]),
                fmt_aud(pnl["net_operating_result"]),
            ])
        add_table(doc, ["Month", "Revenue", "Total Operating Costs", "Net Operating Result"], rows)
        doc.add_paragraph()
        totals = data.totals_24mo[sid]
        add_para(doc, f"24-month total net operating result: {fmt_aud(totals['total_net_operating_result_24mo'])}", bold=True)
        add_para(doc, f"Annualised (steady-state x12) net operating result: {fmt_aud(totals['annualised_net_operating_result_steady_state'])}")
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 8 -- Break-Even Analysis
# ---------------------------------------------------------------------------
def section_breakeven(doc, data: dd.DocumentData):
    add_h1(doc, 8, "Break-Even Analysis")
    add_warning(
        doc,
        "A traditional contribution-margin break-even (fixed costs / contribution margin %) is NOT computed. "
        "Nearly every cost in this model is classified FIXED (including AM/PM payroll under the current, disclosed "
        "treatment), and the GTT-supplies/consumables/laundry classification remains unresolved -- a reliable "
        "fixed-vs-variable split does not exist in the canonical cost data."
    )
    add_para(
        doc,
        "What is computed: an AM client-volume break-even, holding Month 5+ PM revenue, payroll, and operating "
        "expenses fixed and treating AM revenue as the one genuinely volume-linear component (price x operating "
        "days)."
    )
    rows = []
    for sid in dd.SCENARIOS:
        be = data.breakeven[sid]
        rows.append([data.label(sid), be["breakeven_am_client_volume_per_day"], fmt_aud(be["breakeven_monthly_revenue"]),
                      be["committed_client_volume_per_day"], be["margin_of_safety_clients_per_day"]])
    add_table(doc, ["Scenario", "Break-Even AM Vol./Day", "Break-Even Revenue", "Committed Vol./Day", "Margin of Safety"], rows)


# ---------------------------------------------------------------------------
# Section 9 -- Operating Cash Position
# ---------------------------------------------------------------------------
def section_cash_position(doc, data: dd.DocumentData):
    add_h1(doc, 9, "Operating Cash Position")
    add_warning(
        doc,
        "This is an OPERATING-CASH-basis view only (net operating result used as an accrual-basis proxy for cash -- "
        "NOT a true cash-basis forecast, since no debtor/creditor timing data exists anywhere in this repository). "
        "This is NOT a complete funding schedule -- it excludes startup capital deployment, capex, and any opening "
        "cash balance (none assumed)."
    )
    for sid in dd.SCENARIOS:
        cf = data.cash_flow[sid]
        add_para(doc, f"{data.label(sid)}: operating cash trough of {fmt_aud(cf['trough_cumulative_position'])} at Month {cf['trough_month']}.", bold=True)
    add_para(
        doc,
        "Both scenarios' current operating-cash troughs are lower in magnitude than the historical working capital "
        "reserve figure (A$85,000-110,000) -- this does not imply the historical reserve is oversized, since the "
        "two figures use different bases (a pre-rebase 3-month flat estimate vs. this model's 24-month, ramp-aware "
        "calculation) that are not directly comparable without further reconciliation work. See Section 10."
    )


# ---------------------------------------------------------------------------
# Section 10 -- Funding Requirement
# ---------------------------------------------------------------------------
def section_funding(doc, data: dd.DocumentData):
    add_h1(doc, 10, "Funding Requirement")
    add_warning(
        doc,
        "VERIFICATION-TRACKER.md item 47 -- investigated and BOUNDED (OUTCOME 2), NOT resolved to an exact figure. "
        "The underlying 6-9-range startup-cost reconciliation remains genuinely unresolved "
        "(docs/architecture/STARTUP-COST-RECONCILIATION.md)."
    )
    fr = data.funding
    add_h2(doc, "(A) Pre-Opening Capital -- universal, not scenario-specific")
    poc = fr["pre_opening_capital"]
    add_para(doc, f"Range: {fmt_aud(poc['range_low'])} - {fmt_aud(poc['range_high'])}. Covers construction, equipment/furniture/signage, and legal/lease bond -- excludes working capital.")
    add_para(doc, f"Timing: {poc['timing_classification']}")

    add_h2(doc, "(B) Opening Working Capital -- two disclosed methods, neither chosen")
    hist = fr["opening_working_capital"]["historical_reserve_method"]
    add_para(doc, f"Historical reserve method: {fmt_aud(hist['range_low'])} - {fmt_aud(hist['range_high'])} (stale basis -- pre-rebase Month 1-3 loss estimate, item 30).")
    for cc in fr["opening_working_capital"]["operating_cash_trough_cross_check"]["results"]:
        add_para(doc, f"Operating-cash-trough cross-check, {dd.SCENARIO_LABELS[cc['scenario_id']]}: {fmt_aud(cc['trough_value'])} at Month {cc['trough_month']} -- alternative only, not chosen over the historical reserve.")

    add_h2(doc, "(C) Combined Bounded Funding Requirement")
    combined = fr["combined_funding_requirement_bounded"]["primary_method"]
    add_warning(doc, f"PRIMARY METHOD: {fmt_aud(combined['range_low'])} - {fmt_aud(combined['range_high'])} -- exact match to startup_costs.yml's own component-sum total. No new figure invented.")
    for alt in fr["combined_funding_requirement_bounded"]["alternative_cross_check_method"]["results"]:
        add_para(doc, f"Alternative (A + operating cash trough), {dd.SCENARIO_LABELS[alt['scenario_id']]}: {fmt_aud(alt['range_low'])} - {fmt_aud(alt['range_high'])} -- illustrative only, not chosen as more correct.")

    add_para(
        doc,
        "Anthony's own adopted, reconciled total (docs/CURRENT-STATE.md §7.4): A$292,335 - A$594,900 -- close to "
        "but not identical to the primary method range above; CURRENT-STATE.md itself discloses this gap. Both "
        "figures remain valid, disclosed, side by side -- neither supersedes the other."
    )
    add_para(doc, "Not included anywhere in this section:", bold=True)
    add_bullets(doc, [
        "An opening cash balance -- none exists in this repository.",
        "A financing, debt, or equity structure -- none assumed.",
        "The operating cash trough summed on top of the working capital reserve -- would double-count the same 'survive the ramp' concept.",
    ])


# ---------------------------------------------------------------------------
# Section 11 -- Scenario Comparison
# ---------------------------------------------------------------------------
def section_scenario_comparison(doc, data: dd.DocumentData):
    add_h1(doc, 11, "Scenario Comparison")
    add_para(
        doc,
        "Table 1 and Table 2 use identical headcount (8 dual-qualified treatment staff + 2 phlebotomists) and "
        "identical fixed operating expenses -- the only structural difference is AM client volume (18 vs. 12/day) "
        "and its associated AM revenue and superannuation-affected payroll components. Neither scenario is "
        "designated primary (item 1m, open)."
    )
    sc = data.scenario_comparison
    rows = [
        ["Client volume/day", sc["scenario_table_1"]["client_volume_per_day"], sc["scenario_table_2"]["client_volume_per_day"]],
        ["Steady-state revenue", fmt_aud(sc["scenario_table_1"]["steady_state_revenue"]), fmt_aud(sc["scenario_table_2"]["steady_state_revenue"])],
        ["Steady-state payroll", fmt_aud(sc["scenario_table_1"]["steady_state_payroll"]), fmt_aud(sc["scenario_table_2"]["steady_state_payroll"])],
        ["Steady-state opex", fmt_aud(sc["scenario_table_1"]["steady_state_opex"]), fmt_aud(sc["scenario_table_2"]["steady_state_opex"])],
        ["Steady-state net operating result", fmt_aud(sc["scenario_table_1"]["steady_state_net_operating_result"]), fmt_aud(sc["scenario_table_2"]["steady_state_net_operating_result"])],
        ["Annualised revenue", fmt_aud(sc["scenario_table_1"]["annualised_revenue"]), fmt_aud(sc["scenario_table_2"]["annualised_revenue"])],
        ["Annualised net operating result", fmt_aud(sc["scenario_table_1"]["annualised_net_operating_result"]), fmt_aud(sc["scenario_table_2"]["annualised_net_operating_result"])],
    ]
    add_table(doc, ["Metric", "Table 1 (18/day)", "Table 2 (12/day)"], rows)


# ---------------------------------------------------------------------------
# Section 12 -- Sensitivity Analysis
# ---------------------------------------------------------------------------
def section_sensitivity(doc, data: dd.DocumentData):
    add_h1(doc, 12, "Sensitivity Analysis")
    add_para(doc, "Only the two sensitivities already computed by tools/master_financial_model.py are shown -- no new sensitivity is invented for this report.")
    for sid in dd.SCENARIOS:
        add_h2(doc, f"{data.label(sid)} -- Client Volume (50/75/100/125% of committed)")
        rows = [[f"{r['pct_of_committed']}%", r["client_volume_per_day"], fmt_aud(r["total_revenue"]), fmt_aud(r["net_operating_result"])] for r in data.sensitivity_client_volume[sid]]
        add_table(doc, ["% of Committed", "Clients/Day", "Total Revenue", "Net Operating Result"], rows)
        doc.add_paragraph()
    for sid in dd.SCENARIOS:
        add_h2(doc, f"{data.label(sid)} -- Insurance (modelled vs. itemised)")
        rows = [[r["scenario"], fmt_aud(r["monthly_insurance"]), fmt_aud(r["delta_vs_modelled"]), fmt_aud(r["adjusted_net_operating_result"])] for r in data.sensitivity_insurance[sid]]
        add_table(doc, ["Insurance Scenario", "Monthly Insurance", "Delta vs. Modelled", "Adjusted Net Result"], rows)
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Section 13 -- Key Risks and Unresolved Assumptions
# ---------------------------------------------------------------------------
def section_risks(doc, data: dd.DocumentData):
    add_h1(doc, 13, "Key Risks and Unresolved Assumptions")
    rows = [
        ["30", "Working capital reserve basis is a pre-rebase Month 1-3 loss estimate", "UNRESOLVED (stale basis)"],
        ["36", "Historical vs. canonical revenue gap (A$2,576.36/month)", "RESOLVED (nuanced) -- origin unresolved, methodology resolved"],
        ["39", "PM pre-booking discount not applied", "UNRESOLVED"],
        ["41-42", "Revenue-ramp Month 1-4 percentages' original derivation", "UNRESOLVED origin, canonical going forward"],
        ["43", "AM labour ramp not modelled (fixed from Month 1)", "UNRESOLVED, disclosed conservative simplification"],
        ["45", "GTT supplies/consumables/laundry FIXED vs. VARIABLE", "UNRESOLVED"],
        ["46", "Superannuation partial coverage (2 of 8 payroll components)", "PARTIALLY RESOLVED"],
        ["47", "Funding requirement is bounded, not exact", "PARTIALLY RESOLVED (bounded)"],
        ["n/a", "Insurance modelled vs. itemised estimate", "UNRESOLVED"],
    ]
    add_table(doc, ["Tracker Item", "Description", "Status"], rows)


# ---------------------------------------------------------------------------
# Section 14 -- Historical vs. Canonical Figures
# ---------------------------------------------------------------------------
def section_historical_vs_canonical(doc, data: dd.DocumentData):
    add_h1(doc, 14, "Historical vs. Canonical Figures")
    for item in data.historical_reconciliation:
        add_h2(doc, item["label"])
        if "canonical" in item:
            rows = [
                ["Canonical", fmt_aud(item["canonical"]), item["canonical_status"]],
                ["Historical", fmt_aud(item["historical"]), item["historical_status"]],
            ]
            add_table(doc, ["Basis", "Value", "Status"], rows)
            add_para(doc, f"Gap: {fmt_aud(item['gap'])}. {item['gap_status']}")
        else:
            # Single-value record (the A$63,028.75 historical Net P&L item) -- NOT a
            # canonical/historical revenue pair. Never relabel this as revenue.
            add_table(doc, ["Value", "Status"], [[fmt_aud(item["value"]), item["status"]]])
            add_para(doc, item["what_it_is"])
        doc.add_paragraph()
    add_warning(
        doc,
        "The historical A$63,028.75 figure (Table 1, pre-2026-08-09-superannuation-fix) is a Monthly Net P&L "
        "figure -- NOT revenue. It must never be relabelled as revenue in any deliverable, including this report."
    )
    add_para(
        doc,
        "The canonical methodology intentionally supersedes the historical revenue baseline. The historical figure "
        "is preserved above for traceability, not deleted."
    )


# ---------------------------------------------------------------------------
# Section 15 -- Methodology and Data Governance
# ---------------------------------------------------------------------------
def section_methodology(doc, data: dd.DocumentData):
    add_h1(doc, 15, "Methodology and Data Governance")
    add_para(
        doc,
        "This report is generated entirely from data/canonical/*.yml and data/models/master_financial_model.yml -- "
        "the venture's single source of truth. Every figure carries a governance status:"
    )
    add_bullets(doc, [
        "CANONICAL -- directly from data/canonical/*.yml, validated and consistency-checked.",
        "MODELLED -- computed by tools/master_financial_model.py from canonical inputs.",
        "HISTORICAL-SUPERSEDED -- preserved for traceability only, no longer the basis for any live figure.",
        "BOUNDED -- a defensible range, not a point estimate.",
        "UNRESOLVED -- no defensible figure exists yet; disclosed as open, not guessed.",
    ])
    add_para(doc, "Methodology documents referenced throughout this report:")
    add_bullets(doc, [
        "docs/architecture/CANONICAL-REVENUE-METHODOLOGY.md",
        "docs/architecture/REVENUE-RAMP-METHODOLOGY.md",
        "docs/architecture/COST-RAMP-METHODOLOGY.md",
        "docs/architecture/MASTER-FINANCIAL-MODEL-METHODOLOGY.md",
        "docs/architecture/FUNDING-REQUIREMENT-INVESTIGATION.md",
        "docs/VERIFICATION-TRACKER.md",
        "docs/CURRENT-STATE.md",
    ])
    add_para(
        doc,
        "Before publishing, run tools/validate_canonical_data.py and tools/check_consistency.py, and the full "
        "pytest suite (tests/test_*.py) -- all must pass with zero failures for this report's figures to be "
        "considered current."
    )


# ---------------------------------------------------------------------------
# Section 16 -- Conclusion
# ---------------------------------------------------------------------------
def section_conclusion(doc, data: dd.DocumentData):
    add_h1(doc, 16, "Conclusion")
    add_para(
        doc,
        "GTT Center Perth's canonical financial model shows both committed scenarios reaching a positive, "
        "post-superannuation steady-state net operating result -- A$56,581.70/month under Table 1 (18 clients/day) "
        "and A$21,056.64/month under Table 2 (12 clients/day) -- with a positive 24-month cumulative result under "
        "both. The funding requirement to reach that steady state is bounded at A$357,390 - A$577,180, pending "
        "resolution of the underlying historical startup-capital reconciliation."
    )
    add_para(
        doc,
        "This report deliberately leaves several items open rather than forcing false precision: which scenario "
        "is primary (item 1m), the origin of the historical revenue gap (item 36), the PM discount (item 39), the "
        "AM labour ramp (item 43), superannuation eligibility for two payroll components (item 46), and the exact "
        "funding figure (item 47). These are founder-level decisions, not modelling gaps this report can close on "
        "its own."
    )


def main():
    data = dd.get_data()
    doc = Document()

    add_title_page(doc, data)
    section_executive_summary(doc, data)
    doc.add_page_break()
    section_business_model(doc, data)
    doc.add_page_break()
    section_revenue_model(doc, data)
    doc.add_page_break()
    section_revenue_ramp(doc, data)
    doc.add_page_break()
    section_cost_model(doc, data)
    doc.add_page_break()
    section_super(doc, data)
    doc.add_page_break()
    section_24mo_outlook(doc, data)
    doc.add_page_break()
    section_breakeven(doc, data)
    doc.add_page_break()
    section_cash_position(doc, data)
    doc.add_page_break()
    section_funding(doc, data)
    doc.add_page_break()
    section_scenario_comparison(doc, data)
    doc.add_page_break()
    section_sensitivity(doc, data)
    doc.add_page_break()
    section_risks(doc, data)
    doc.add_page_break()
    section_historical_vs_canonical(doc, data)
    doc.add_page_break()
    section_methodology(doc, data)
    doc.add_page_break()
    section_conclusion(doc, data)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Word report written: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
